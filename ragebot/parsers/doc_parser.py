"""
Document Parser — Extracts, chunks, and *summarises* content from
PDF, DOCX, MD, and TXT files.

Summarisation strategy (in priority order):
  1. NLTK extractive summariser (sentence scoring by TF-IDF-ish frequency)
  2. Simple first-paragraph fallback
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Any


import logging

# Suppress noisy HuggingFace logs
from ragebot.utils.logging_config import suppress_noisy_logs
suppress_noisy_logs()

# ── NLTK Summariser ───────────────────────────────────────────────────────────

def _ensure_nltk_data() -> bool:
    """Download required NLTK data on first use.  Returns True if NLTK is usable."""
    try:
        import nltk  # type: ignore
        for resource in ("punkt", "punkt_tab", "stopwords"):
            try:
                nltk.data.find(f"tokenizers/{resource}" if "punkt" in resource else f"corpora/{resource}")
            except LookupError:
                nltk.download(resource, quiet=True)
        return True
    except Exception:
        return False


def _nltk_summarise(text: str, sentence_count: int = 5) -> str:
    """Extractive summary using NLTK sentence + word frequency scoring."""
    try:
        import nltk  # type: ignore
        from nltk.corpus import stopwords  # type: ignore
        from nltk.tokenize import sent_tokenize, word_tokenize  # type: ignore

        if not _ensure_nltk_data():
            return ""

        stops = set(stopwords.words("english"))
        sentences = sent_tokenize(text)
        if len(sentences) <= sentence_count:
            return text[:1500]

        # Word frequency table (excluding stopwords)
        word_freq: dict[str, int] = {}
        for word in word_tokenize(text.lower()):
            if word.isalnum() and word not in stops:
                word_freq[word] = word_freq.get(word, 0) + 1

        if not word_freq:
            return text[:1500]

        max_freq = max(word_freq.values())
        for w in word_freq:
            word_freq[w] /= max_freq  # normalise 0-1

        # Score each sentence
        scored: list[tuple[float, int, str]] = []
        for idx, sent in enumerate(sentences):
            score = sum(word_freq.get(w, 0) for w in word_tokenize(sent.lower()) if w.isalnum())
            scored.append((score, idx, sent))

        scored.sort(key=lambda x: x[0], reverse=True)
        top = sorted(scored[:sentence_count], key=lambda x: x[1])  # restore doc order
        return " ".join(s for _, _, s in top)
    except Exception:
        return ""


# ── Document Parser ───────────────────────────────────────────────────────────

class DocumentParser:
    def parse(self, file_path: Path) -> dict:
        """Parse a document file and return structured content."""
        ext = file_path.suffix.lower()
        parsers = {
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".md": self._parse_markdown,
            ".txt": self._parse_text,
            ".rst": self._parse_text,
            ".tex": self._parse_text,
        }
        parser_fn = parsers.get(ext, self._parse_text)
        return parser_fn(file_path)

    # ── Parsers ───────────────────────────────────────────────────────────────

    def _parse_pdf(self, path: Path) -> dict:
        text = ""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(path))
            pages = []
            for page in doc:
                pages.append(page.get_text())
            text = "\n".join(pages)
            doc.close()
        except ImportError:
            try:
                import pdfplumber
                with pdfplumber.open(str(path)) as pdf:
                    text = "\n".join(p.extract_text() or "" for p in pdf.pages)
            except ImportError:
                text = "[PDF parsing unavailable. Install PyMuPDF: pip install pymupdf]"
        except Exception as e:
            text = f"[PDF parsing error: {e}]"

        return self._build_result(text, "pdf", str(path))

    def _parse_docx(self, path: Path) -> dict:
        text = ""
        try:
            from docx import Document
            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            text = "\n".join(paragraphs)
        except ImportError:
            text = "[DOCX parsing unavailable. Install python-docx: pip install python-docx]"
        except Exception as e:
            text = f"[DOCX parsing error: {e}]"

        return self._build_result(text, "docx", str(path))

    def _parse_markdown(self, path: Path) -> dict:
        try:
            text = path.read_text(encoding="utf-8", errors="ignore")
        except Exception as e:
            return self._build_result("", "markdown", str(path))

        # Extract headings for structure
        headings = re.findall(r"^#{1,6}\s+(.+)$", text, re.MULTILINE)

        # Strip markdown syntax for clean text
        clean = re.sub(r"```[\s\S]*?```", "[code block]", text)
        clean = re.sub(r"`[^`]+`", lambda m: m.group()[1:-1], clean)
        clean = re.sub(r"!\[.*?\]\(.*?\)", "[image]", clean)
        clean = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", clean)
        clean = re.sub(r"#{1,6}\s+", "", clean)
        clean = re.sub(r"\*\*(.+?)\*\*", r"\1", clean)
        clean = re.sub(r"\*(.+?)\*", r"\1", clean)

        result = self._build_result(clean, "markdown", str(path))
        result["headings"] = headings
        return result

    def _parse_text(self, path: Path) -> dict:
        try:
            text = path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            text = ""
        return self._build_result(text, "text", str(path))

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _build_result(self, text: str, doc_type: str, file_path: str) -> dict:
        text = text.strip()
        summary = self._extract_summary(text)
        chunks = self._chunk_text(text)
        return {
            "type": doc_type,
            "summary": summary,
            "chunks": chunks,
            "char_count": len(text),
            "chunk_count": len(chunks),
        }

    def _extract_summary(self, text: str, max_chars: int = 500) -> str:
        """
        Produce a summary of the document text.

        Strategy:
          1. Try NLTK extractive summariser (picks the most informative sentences)
          2. Fall back to first meaningful paragraph
        """
        if not text:
            return ""

        # ── Attempt NLTK ──────────────────────────────────────────────────
        nltk_summary = _nltk_summarise(text, sentence_count=5)
        if nltk_summary and len(nltk_summary) > 50:
            return nltk_summary[:max_chars]

        # ── Fallback: first paragraph ─────────────────────────────────────
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        for para in paragraphs:
            if len(para) > 50:
                return para[:max_chars]
        return text[:max_chars]

    def _chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
        """Split text into overlapping chunks."""
        if not text:
            return []
        words = text.split()
        if not words:
            return []

        chunks = []
        start = 0
        while start < len(words):
            end = start + chunk_size
            chunk_words = words[start:end]
            chunks.append(" ".join(chunk_words))
            if end >= len(words):
                break
            start = end - overlap  # overlap

        return chunks if chunks else [text[:2000]]
